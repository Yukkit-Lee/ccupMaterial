from pathlib import Path
import subprocess
import zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "result" / "智检云盾商业计划书（写作指导重写版）.md"
WORK_MD = ROOT / "result" / "_build_word.md"
OUTPUT = ROOT / "result" / "智检云盾商业计划书（图表修复版）.docx"

text = SOURCE.read_text(encoding="utf-8")
text = text.replace("charts/delivery-flow.png", "word_charts/delivery-flow.jpg")
text = text.replace("charts/customer-plan.png", "word_charts/customer-plan.jpg")
text = text.replace("charts/revenue-composition.png", "word_charts/revenue-composition.jpg")
text = text.replace("charts/financial-forecast.png", "word_charts/financial-forecast.jpg")
text = text.replace("charts/cash-use.png", "word_charts/cash-use.jpg")
WORK_MD.write_text(text, encoding="utf-8")

pandoc = Path(r"C:\Program Files\Pandoc\pandoc.exe")
command = [
    str(pandoc), str(WORK_MD),
    "--from", "gfm+pipe_tables",
    "--resource-path", str(ROOT / "result"),
    "--output", str(OUTPUT),
]
subprocess.run(command, check=True, cwd=ROOT)

# Make the default Chinese font explicit in the generated document and keep all
# figures inline, which is the most reliable placement mode across Word/LibreOffice.
doc = Document(OUTPUT)
for section in doc.sections:
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for style_name, size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 14)]:
    style = doc.styles[style_name]
    style.font.name = "黑体"
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
for para in doc.paragraphs:
    for run in para.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for shape in doc.inline_shapes:
    # 6.0 inches ensures the figure remains inside the planned text area.
    if shape.width > 5486400:
        ratio = 5486400 / shape.width
        shape.width = 5486400
        shape.height = int(shape.height * ratio)
for para in doc.paragraphs:
    if para.text.startswith("图 "):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(11)
doc.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    media = [n for n in archive.namelist() if n.startswith("word/media/")]
    embedded_figures = [n for n in media if n.lower().endswith(".jpg")]
    document_xml = archive.read("word/document.xml")
    if len(media) != 5 or len(embedded_figures) != 5 or b"r:link=" in document_xml:
        raise RuntimeError(f"Expected five embedded, non-external JPEG figures; got {media}")
print(f"Created {OUTPUT.name} with {len(embedded_figures)} embedded JPEG figures.")
