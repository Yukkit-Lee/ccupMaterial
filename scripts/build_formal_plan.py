from pathlib import Path
import zipfile
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "result" / "智检云盾商业计划书（正式提交版）.docx"
CHARTS = ROOT / "result" / "word_charts"
RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x79)

def set_font(run, size=12, bold=False, color=BLACK, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = int(sum(widths_cm) / 2.54 * 1440)
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    layout = tblPr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def add_run(paragraph, value, red=False, size=12, bold=False):
    run = paragraph.add_run(value)
    set_font(run, size=size, bold=bold, color=RED if red else BLACK)
    return run

def add_para(doc, text="", red=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, after=5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.74) if align == WD_ALIGN_PARAGRAPH.JUSTIFY else Cm(0)
    p.paragraph_format.line_spacing = Pt(25)
    p.paragraph_format.space_after = Pt(after)
    add_run(p, text, red=red, size=size, bold=bold)
    return p

def add_mixed(doc, chunks, after=5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(25)
    p.paragraph_format.space_after = Pt(after)
    for value, red, bold in chunks:
        add_run(p, value, red=red, bold=bold)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else BLACK, name="黑体")
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = Pt(23)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, item, red=item.startswith("【"), size=11.5)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, "D9EAF7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, value, red="【此处填写" in value, size=10.2, bold=value.startswith("合计"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_figure(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(CHARTS / filename), width=Cm(15.0))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(7)
    add_run(c, caption, size=10.5)

def page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    set_font(run, size=9)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3.5)
section.bottom_margin = Cm(3.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)
section.header_distance = Cm(1.5)
section.footer_distance = Cm(1.5)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(header, "智检云盾商业计划书（正式提交版）", size=9)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer, "第 ", size=9); page_field(footer); add_run(footer, " 页", size=9)

# 封面：proposal_centerpiece 样式。
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "第十五届“挑战杯”中国大学生创业计划竞赛", size=16, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "东北振兴产业升级专项赛·数智赋能产业", size=13)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "智检云盾", size=28, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "——面向装备制造的多工厂数据安全协同质检平台", size=18, bold=True)
doc.add_paragraph()
meta = doc.add_table(rows=4, cols=2)
set_table_widths(meta, [3.2, 12.2])
for row, (label, value, is_red) in zip(meta.rows, [
    ("项目负责人", "李旭杰", False),
    ("工业应用与产品运营负责人", "胡思佳", False),
    ("申报单位", "【此处填写申报单位名称】", True),
    ("项目联系人", "【此处填写联系人、电话、邮箱及日期】", True),
]):
    shade(row.cells[0], "D9EAF7")
    p1 = row.cells[0].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_run(p1, label, size=11, bold=True)
    p2 = row.cells[1].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.LEFT; add_run(p2, value, red=is_red, size=11)
doc.add_page_break()

add_heading(doc, "一、执行摘要")
add_para(doc, "智检云盾是面向装备制造产业链的多工厂数据安全协同质检平台。平台针对划伤、裂纹、凹坑、气孔、焊缝和涂装异常等视觉质检任务，解决罕见缺陷样本分散、企业之间难以共享原始图像与工艺数据的问题。参与企业在本地保留图像、批次与工艺信息，平台在约定任务范围内组织安全协同训练，使模型能够利用多方经验而不直接汇集原始数据。")
add_para(doc, "平台以“本地数据治理—安全协同训练—改进 DPSGD 隐私保护更新—本地部署与人工复检”为闭环，服务汽车零部件、重大装备、电站装备和轨道交通装备等场景。首期聚焦二维表面缺陷分类，通过 POC 验证数据边界、标签规则、模型效果与现场流程，再逐步进入私有化部署和订阅运维。")
add_mixed(doc, [("当前原型与测试进展：", False, True), ("【此处填写已完成的原型功能、测试数据、实验结果及证明材料编号】", True, False)])

add_heading(doc, "二、项目概述")
add_heading(doc, "1. 场景痛点", 2)
add_para(doc, "单厂视觉质检模型容易受长尾缺陷样本不足、标签口径不一致和工况差异影响；而集中采集供应链图像、工艺参数和批次信息，又会触及企业工艺秘密、质量责任和数据权限边界。智检云盾将协同学习限定在经授权的任务、参与方和模型版本内，并保留训练、权限与模型的审计记录。")
add_heading(doc, "2. 产品构成", 2)
add_bullets(doc, [
    "本地数据治理模块：标签校验、数据版本、最小必要采集与访问权限管理。",
    "安全协同训练模块：秘密共享训练、梯度裁剪、改进 DPSGD 隐私保护更新与训练编排。",
    "质检应用模块：本地推理、缺陷类别与置信度输出、人工复检队列和结果回写。",
    "审计运维模块：任务授权、参与方权限、模型版本、隐私预算、日志追溯与回滚。",
])
add_figure(doc, "delivery-flow.jpg", "图 2.1  智检云盾首期交付闭环")

add_heading(doc, "三、项目优势")
add_para(doc, "第一，数据不出厂。原始质检图像、工艺参数和批次信息保留在数据持有方环境中，协同训练只处理约定范围内的受保护信息。第二，场景化交付。平台将缺陷类别、标签口径、人工复检、模型版本和验收流程组织为统一场景包，避免通用技术与现场质量流程脱节。第三，自主研发的改进 DPSGD 算法与安全协同训练链路面向少量合作主体和轻量缺陷识别任务进行适配，兼顾隐私保护、通信开销与训练可用性。")
add_mixed(doc, [("技术优势的量化证明：", False, True), ("【此处填写改进点、对比基线、样本量、准确率/召回率、时延、通信量、隐私预算及实验记录编号】", True, False)])

add_heading(doc, "四、市场分析与客户价值")
add_para(doc, "首期目标客户为具有明确视觉质检任务、至少两方可协作数据持有主体、对工艺和质量数据有保密要求、并愿意通过小范围 POC 验证价值的装备制造企业。质量负责人关注漏检、误报和复检效率，信息化负责人关注节点安全和运维，供应链质量负责人关注标签一致性与责任边界，企业管理者关注投入节奏和验收结果。")
add_para(doc, "项目以汽车零部件表面缺陷为首期场景，后续可复制到重大装备、轨道交通和电站装备的相似视觉任务。市场进入采取“场景成熟度优先”的策略：先选择已有视觉采集条件、可形成复检口径、能够明确数据负责人和验收责任人的场景，再通过检测机构、智能制造集成商、产业园与行业协会拓展触达。")
add_mixed(doc, [("客户访谈与试点证明：", False, True), ("【此处填写访谈对象、日期、核心需求、意向进度、授权情况及附件编号】", True, False)])

add_heading(doc, "五、商业模式")
add_para(doc, "项目采用“POC 验证—标准部署—年度订阅—定制集成”的分阶段商业模式。POC 用于确认数据边界、缺陷标签与验收指标；标准部署完成本地节点、模型与审计配置；年度订阅覆盖模型迭代、运行审计与支持；定制集成适配现场接口和质量台账。合同应以需求确认、阶段交付和验收结果为基础设置回款节点。")
add_table(doc, ["产品/服务", "拟定报价", "主要交付内容", "直接成本假设"], [
    ("场景 POC", "6 万元/场景", "数据边界、三方训练、评测与验收报告", "2.5 万元/场景"),
    ("标准私有化部署", "18 万元/节点", "节点安装、模型交付、权限与审计配置", "8.5 万元/节点"),
    ("年度订阅服务", "8 万元/节点/年", "版本维护、再训练、运行审计与支持", "2.5 万元/节点/年"),
    ("定制集成", "10 万元/项", "接口、台账或流程适配", "4 万元/项"),
], [3.0, 2.4, 6.0, 4.0])
add_mixed(doc, [("报价与成本的调研依据：", False, True), ("【此处填写实际报价来源、工时、差旅、算力、税费及报价日期】", True, False)])

add_heading(doc, "六、团队与组织")
add_table(doc, ["成员", "角色", "职责", "需补充材料"], [
    ("李旭杰", "项目负责人、算法负责人", "总体推进、算法路线、安全协同训练原型、模型评测与合规统筹", "【此处填写专业/学历、项目经历、联系方式及证明材料】"),
    ("胡思佳", "工业应用负责人、产品运营负责人", "质检流程梳理、标签规范、POC 验收、用户调研与运营材料", "【此处填写专业/学历、项目经历、联系方式及证明材料】"),
    ("【此处填写成员姓名】", "【此处填写角色】", "【此处填写实际职责】", "【此处填写专业/经历/贡献】"),
], [2.1, 3.0, 5.5, 4.8])
add_para(doc, "项目现有登记主体为永清县臻青奇百货店（个体工商户），统一社会信用代码 92131023MAEB0MJR4L，经营者为李旭杰，资金数额 1 万元，成立日期为 2025 年 1 月 18 日，当前状态为存续。")
add_mixed(doc, [("主体承接与资质说明：", False, True), ("【此处填写营业执照经营范围、项目承接关系、开票/合同/数据处理安排及证明材料编号】", True, False)])

add_heading(doc, "七、财务分析")
add_para(doc, "以下为项目经营预测，收入由服务单价与客户数量逐项推导。首年以 POC 和首个部署为主，处于研发与试点投入期；第二年依托部署和订阅逐步改善经营结果；第三年在标准化交付、续费和定制集成的共同支撑下形成规模化收入。")
add_table(doc, ["项目（万元）", "2027E", "2028E", "2029E"], [
    ("营业收入", "38", "156", "326"),
    ("直接交付成本", "16", "65", "133"),
    ("毛利", "22", "91", "193"),
    ("研发与交付人力", "38", "50", "75"),
    ("销售与客户实施", "6", "12", "18"),
    ("管理、合规及云资源", "9", "14", "20"),
    ("税前经营性利润", "-31", "15", "80"),
], [6.2, 3.0, 3.0, 3.0])
add_figure(doc, "customer-plan.jpg", "图 7.1  三年客户服务数量计划")
add_figure(doc, "revenue-composition.jpg", "图 7.2  三年营业收入构成")
add_figure(doc, "financial-forecast.jpg", "图 7.3  营业收入与税前经营性利润预测")
add_mixed(doc, [("财务预测的依据与修订口径：", False, True), ("【此处填写客户漏斗、转化率、续费率、回款周期、成本测算表及相关附件编号】", True, False)])

add_heading(doc, "八、资金规划")
add_para(doc, "项目拟将资金优先用于原型研发与安全测试、POC 交付与现场适配、数据治理与合规、市场调研与渠道建设，以及必要的流动资金缓冲。资金使用应与阶段性里程碑和回款安排匹配，避免在首期购置与场景验证无关的重资产。")
add_table(doc, ["资金用途", "计划金额（万元）", "说明"], [
    ("原型研发与安全测试", "30", "训练编排、审计与测试工具"),
    ("POC 交付与现场适配", "20", "试点实施、调研与节点适配"),
    ("数据治理与合规", "12", "数据处理约定、标签规范与测试"),
    ("市场调研与渠道", "8", "客户访谈、伙伴拓展与行业活动"),
    ("流动资金缓冲", "10", "应对验收回款周期"),
    ("合计", "80", "启动资金池预算"),
], [5.0, 3.8, 6.4])
add_figure(doc, "cash-use.jpg", "图 8.1  启动资金用途与现金缓冲安排")
add_mixed(doc, [("现有资金与融资安排：", False, True), ("【此处填写实际可投入资金、已发生支出、融资需求、融资方式、资金来源及证明材料】", True, False)])

add_heading(doc, "九、风险控制")
add_table(doc, ["风险类别", "主要表现", "应对措施"], [
    ("技术与精度", "隐私保护与安全计算可能影响模型精度和训练时延", "先在三方、图像分类和浅层模型范围内验证；以准确率、时延、通信量和隐私预算联合验收"),
    ("数据与合规", "工艺秘密、批次信息或客户数据处理不当", "数据不出厂、最小必要采集、权限分级、日志审计与退出删除机制"),
    ("市场与回款", "客户决策慢、POC 不转部署、验收延迟", "分阶段验收与回款；控制固定成本；优先选择场景成熟度高的客户"),
    ("主体与知识产权", "主体资质、代码来源或组件授权不清", "核验经营范围；建立代码、数据和组件来源清单；保留研发记录与许可证"),
], [3.0, 5.4, 6.8])

add_heading(doc, "十、发展规划")
add_para(doc, "近期完成三方模拟原型与一类表面缺陷场景包，形成可复现实验、标签规范、训练记录和 POC 验收模板；中期以汽车零部件场景为样板，复制标准部署和订阅运维能力；远期在客户授权和真实案例基础上，扩展至重大装备、轨道交通和电站装备的相似视觉任务。")
add_mixed(doc, [("里程碑与责任人：", False, True), ("【此处填写各阶段起止时间、交付物、责任人、验收指标及真实资源保障】", True, False)])

add_heading(doc, "附件目录")
add_bullets(doc, [
    "【此处填写营业执照、主体关系说明及团队成员证明材料】",
    "【此处填写原型截图、代码说明、实验记录、数据授权或公开数据来源说明】",
    "【此处填写用户访谈、意向函、保密协议、试点协议或授权材料】",
    "【此处填写报价测算、成本明细、融资/资金来源及其他支撑材料】",
])

doc.save(OUT)
with zipfile.ZipFile(OUT) as archive:
    media = [n for n in archive.namelist() if n.startswith("word/media/")]
    document_xml = archive.read("word/document.xml")
    if len(media) != 5 or b"r:link=" in document_xml:
        raise RuntimeError("Formal plan must contain five embedded, non-linked figures.")
print(f"Created {OUT.name} with {len(media)} embedded figures.")
