from pathlib import Path
import subprocess
import zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "result" / "智检云盾商业计划书（写作指导重写版）.md"
FORMAL_MD = ROOT / "result" / "智检云盾商业计划书（正式提交版·完整）.md"
WORK_MD = ROOT / "result" / "_formal_full_build.md"
OUTPUT = ROOT / "result" / "智检云盾商业计划书（正式提交版·完整修订版）.docx"

source = SOURCE.read_text(encoding="utf-8")
start = source.index("# 一、执行摘要")
end = source.index("# 参考资料与提交前核验事项")
body = source[start:end]

# 删除仅面向写作过程的说明，保留项目、技术、市场、商业与财务正文。
body = body.replace(
    "由于尚未提供指导教师信息，本计划书不列示具体导师；提交竞赛材料时也应按要求隐去推报单位和导师身份信息。",
    ""
)
body = body.replace(
    "竞赛提交版还应按要求隐去指导教师姓名等推报单位信息。",
    ""
)
team_marker = "## 2. 顾问与协作原则"
team_fill = """## 1.1 团队基本信息

李旭杰，项目负责人兼算法负责人，海南大学网络空间安全学院硕士研究生（2025 年至今）。研究方向聚焦隐私计算与人工智能安全，系统学习现代密码学、机器学习、安全多方计算和云计算安全等相关理论与方法，持续跟踪可搜索加密、云计算安全与安全多方计算领域研究进展。在本项目中负责总体技术路线与算法方案设计，围绕多工厂数据安全协同质检场景开展改进 DPSGD 算法、安全协同训练原型、模型评测与隐私预算配置等工作。

【此处填写胡思佳的学校、专业、年级或学历、个人简介及联系方式】

【此处填写其他团队成员姓名、学校/专业、职责及个人简介；如无其他成员请删除本行】

叶俊于中国西安西安电子科技大学获得博士学位，现为海南大学网络空间安全学院教授、博士生导师。他是一名机器学习工程师、人工智能应用工程师和密码学安全工程师。其研究兴趣包括应用密码学、云计算和人工智能。他已在国内外知名学术期刊和会议上发表或合作发表 100 余篇研究论文，并担任多家国际期刊的客座编辑和审稿人，以及众多国内外会议的程序委员会委员。

"""
body = body.replace(team_marker, team_fill + team_marker)

front = """# 第十五届“挑战杯”中国大学生创业计划竞赛商业计划书

## 智检云盾——面向装备制造的多工厂数据安全协同质检平台

赛道：东北振兴产业升级专项赛·数智赋能产业  
申报主体：永清县臻青奇百货店（个体工商户）  
项目负责人、算法负责人：李旭杰  
工业应用负责人、产品运营负责人：胡思佳

"""

formal = front + body
FORMAL_MD.write_text(formal, encoding="utf-8")
word_text = formal.replace("charts/delivery-flow.png", "word_charts/delivery-flow.jpg")
word_text = word_text.replace("charts/customer-plan.png", "word_charts/customer-plan.jpg")
word_text = word_text.replace("charts/revenue-composition.png", "word_charts/revenue-composition.jpg")
word_text = word_text.replace("charts/financial-forecast.png", "word_charts/financial-forecast.jpg")
word_text = word_text.replace("charts/cash-use.png", "word_charts/cash-use.jpg")
WORK_MD.write_text(word_text, encoding="utf-8")

pandoc = Path(r"C:\Program Files\Pandoc\pandoc.exe")
subprocess.run([
    str(pandoc), str(WORK_MD),
    "--from", "gfm+pipe_tables",
    "--resource-path", str(ROOT / "result"),
    "--output", str(OUTPUT),
], check=True, cwd=ROOT)

doc = Document(OUTPUT)
for section in doc.sections:
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for name, size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
    style = doc.styles[name]
    style.font.name = "黑体"
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

for para in doc.paragraphs:
    if para.text.startswith("图 "):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if "【此处填写" in run.text:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.bold = True
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    if "【此处填写" in run.text:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        run.bold = True
for shape in doc.inline_shapes:
    if shape.width > 5486400:
        scale = 5486400 / shape.width
        shape.width = 5486400
        shape.height = int(shape.height * scale)
doc.save(OUTPUT)

with zipfile.ZipFile(OUTPUT) as archive:
    media = [n for n in archive.namelist() if n.startswith("word/media/")]
    document_xml = archive.read("word/document.xml")
    if len(media) != 5 or b"r:link=" in document_xml:
        raise RuntimeError("Expected five embedded, non-linked figures.")
print(f"Created {OUTPUT.name}; source length: {len(formal)} characters; embedded figures: {len(media)}.")
